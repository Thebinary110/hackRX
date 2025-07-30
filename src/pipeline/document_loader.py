import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import email
from bs4 import BeautifulSoup
import chardet
from typing import Union
from io import BytesIO
import requests
import tempfile
import os
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try different PDF libraries with better error handling
PYMUPDF_AVAILABLE = False
PYPDF2_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    # Test if fitz.open is available
    if hasattr(fitz, 'open'):
        PYMUPDF_AVAILABLE = True
        logger.info(f"✅ PyMuPDF available (version: {getattr(fitz, '__version__', 'Unknown')})")
    else:
        logger.warning("⚠️ PyMuPDF imported but 'open' method not available")
except ImportError as e:
    logger.warning(f"⚠️ PyMuPDF not available: {e}")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
    logger.info(f"✅ PyPDF2 available (version: {getattr(PyPDF2, '__version__', 'Unknown')})")
except ImportError as e:
    logger.warning(f"⚠️ PyPDF2 not available: {e}")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    logger.info(f"✅ pdfplumber available")
except ImportError as e:
    logger.warning(f"⚠️ pdfplumber not available: {e}")

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
    logger.info("✅ OCR libraries available")
except ImportError as e:
    OCR_AVAILABLE = False
    logger.warning(f"⚠️ OCR libraries not available: {e}")

def load_and_clean(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from various file formats.
    
    Args:
        file_bytes: The file content as bytes
        filename: The filename or URL
        
    Returns:
        Extracted text as string
    """
    logger.info(f"Processing file: {filename}")
    
    try:
        if filename.startswith("http"):  # web URL case
            text = extract_text_from_url(filename)
        elif filename.lower().endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        elif filename.lower().endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif filename.lower().endswith((".eml", ".msg")):
            text = extract_text_from_email(file_bytes)
        elif filename.lower().endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
            text = extract_text_from_image(file_bytes)
        else:
            # Try to detect if it's a text file
            try:
                encoding = chardet.detect(file_bytes)["encoding"] or "utf-8"
                text = file_bytes.decode(encoding, errors="ignore")
            except:
                raise ValueError(f"Unsupported file format: {filename}")
        
        if not text or not text.strip():
            raise ValueError(f"No text could be extracted from {filename}")
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error processing file {filename}: {str(e)}")
        raise ValueError(f"Error processing file {filename}: {str(e)}")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using available PDF libraries with improved error handling.
    """
    if not any([PYMUPDF_AVAILABLE, PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE]):
        raise ValueError("No PDF processing library available. Please install PyMuPDF, PyPDF2, or pdfplumber using:\npip install PyMuPDF PyPDF2 pdfplumber")
    
    # Method 1: PyMuPDF (fastest and most reliable)
    if PYMUPDF_AVAILABLE:
        try:
            logger.info("Trying PyMuPDF...")
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            page_count = doc.page_count
            logger.info(f"PDF has {page_count} pages")
            
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    logger.debug(f"Page {page_num + 1}: {len(page_text)} characters")
                else:
                    logger.debug(f"Page {page_num + 1}: No text found")
            
            doc.close()
            
            if text.strip():
                logger.info(f"✅ PyMuPDF extracted {len(text)} characters from {page_count} pages")
                return text.strip()
            else:
                logger.warning("PyMuPDF: No text extracted (possibly scanned/image PDF)")
                
        except Exception as e:
            logger.error(f"PyMuPDF failed: {e}")
    
    # Method 2: PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            logger.info("Trying PyPDF2...")
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text = ""
            page_count = len(reader.pages)
            logger.info(f"PDF has {page_count} pages")
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    logger.debug(f"Page {page_num + 1}: {len(page_text)} characters")
                else:
                    logger.debug(f"Page {page_num + 1}: No text found")
            
            if text.strip():
                logger.info(f"✅ PyPDF2 extracted {len(text)} characters from {page_count} pages")
                return text.strip()
            else:
                logger.warning("PyPDF2: No text extracted")
                
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
    
    # Method 3: pdfplumber
    if PDFPLUMBER_AVAILABLE:
        try:
            logger.info("Trying pdfplumber...")
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                text = ""
                page_count = len(pdf.pages)
                logger.info(f"PDF has {page_count} pages")
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        logger.debug(f"Page {page_num + 1}: {len(page_text)} characters")
                    else:
                        logger.debug(f"Page {page_num + 1}: No text found")
                
                if text.strip():
                    logger.info(f"✅ pdfplumber extracted {len(text)} characters from {page_count} pages")
                    return text.strip()
                else:
                    logger.warning("pdfplumber: No text extracted")
                    
        except Exception as e:
            logger.error(f"pdfplumber failed: {e}")
    
    # If we get here, all methods failed
    error_msg = "All PDF processing methods failed. "
    if not text or not text.strip():
        error_msg += "This might be a scanned PDF or image-based PDF that requires OCR processing."
    
    raise ValueError(error_msg)

def extract_text_from_pdf_with_ocr(file_bytes: bytes) -> str:
    """
    Extract text from PDF using OCR (for scanned/image PDFs).
    This is a fallback method when regular text extraction fails.
    """
    if not OCR_AVAILABLE:
        raise ValueError("OCR libraries not available. Install with: pip install pytesseract pillow")
    
    if not PYMUPDF_AVAILABLE:
        raise ValueError("PyMuPDF required for OCR extraction")
    
    try:
        logger.info("Attempting OCR extraction...")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            
            # OCR the image
            image = Image.open(BytesIO(img_data))
            page_text = pytesseract.image_to_string(image)
            
            if page_text.strip():
                text += page_text + "\n"
                logger.debug(f"OCR Page {page_num + 1}: {len(page_text)} characters")
        
        doc.close()
        
        if text.strip():
            logger.info(f"✅ OCR extracted {len(text)} characters")
            return text.strip()
        else:
            raise ValueError("OCR could not extract any text")
            
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise ValueError(f"OCR extraction failed: {e}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Enhanced DOCX text extraction that handles tables, headers, footers."""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        full_text = []
        
        # Method 1: Extract from paragraphs and tables in document order
        def iter_block_items(parent):
            if isinstance(parent, Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("something's not right")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        # Extract text in document order
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                para_text = block.text.strip()
                if para_text:
                    full_text.append(para_text)
            elif isinstance(block, Table):
                table_text = extract_table_text(block)
                if table_text:
                    full_text.append(table_text)
        
        # Extract from headers and footers
        for section in doc.sections:
            # Headers
            header = section.header
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[HEADER] {para.text.strip()}")
            
            # Footers
            footer = section.footer
            if footer:
                for para in footer.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[FOOTER] {para.text.strip()}")
        
        result = '\n\n'.join(full_text)
        
        if not result.strip():
            raise ValueError("No text could be extracted from the DOCX file")
        
        return result
        
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")

def extract_table_text(table) -> str:
    """Extract text from a DOCX table."""
    table_data = []
    try:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            if row_data:
                table_data.append(' | '.join(row_data))
        
        if table_data:
            return '\n'.join(table_data)
    except Exception as e:
        logger.error(f"Error extracting table text: {e}")
    
    return ""

def extract_text_from_email(file_bytes: bytes) -> str:
    """Extract text from email files (.eml, .msg)."""
    try:
        encoding_info = chardet.detect(file_bytes)
        encoding = encoding_info.get("encoding", "utf-8")
        
        msg = email.message_from_bytes(file_bytes)
        body_parts = []
        
        # Extract subject
        subject = msg.get("Subject", "")
        if subject:
            body_parts.append(f"Subject: {subject}")
        
        # Extract sender and recipient info
        sender = msg.get("From", "")
        if sender:
            body_parts.append(f"From: {sender}")
        
        recipient = msg.get("To", "")
        if recipient:
            body_parts.append(f"To: {recipient}")
        
        body_parts.append("")  # Empty line separator
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        text = payload.decode(encoding, errors="ignore").strip()
                        if text:
                            body_parts.append(text)
                elif content_type == "text/html":
                    payload = part.get_payload(decode=True)
                    if payload:
                        html = payload.decode(encoding, errors="ignore")
                        soup = BeautifulSoup(html, "html.parser")
                        text = soup.get_text(separator="\n").strip()
                        if text:
                            body_parts.append(text)
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                content_type = msg.get_content_type()
                if content_type == "text/html":
                    soup = BeautifulSoup(payload.decode(encoding, errors="ignore"), "html.parser")
                    text = soup.get_text(separator="\n").strip()
                else:
                    text = payload.decode(encoding, errors="ignore").strip()
                
                if text:
                    body_parts.append(text)
        
        result = '\n'.join(body_parts)
        if not result.strip():
            raise ValueError("No content could be extracted from email")
        
        return result
        
    except Exception as e:
        raise ValueError(f"Error extracting text from email: {str(e)}")

def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from image using OCR."""
    if not OCR_AVAILABLE:
        raise ValueError("OCR libraries (pytesseract, PIL) not available. Please install them.")
    
    try:
        image = Image.open(BytesIO(file_bytes))
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        text = pytesseract.image_to_string(image, config='--psm 6')
        if not text.strip():
            raise ValueError("No text could be extracted from image")
        
        return text.strip()
        
    except Exception as e:
        raise ValueError(f"Error extracting text from image: {str(e)}")

def extract_text_from_url(url: str) -> str:
    """Extract text from web URL."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Handle blob URLs or special URLs
        if 'blob.core.windows.net' in url or 'extension://' in url:
            # Extract the actual PDF URL
            if '?file=' in url:
                import urllib.parse
                actual_url = urllib.parse.unquote(url.split('?file=')[1])
                if actual_url.startswith('https://'):
                    url = actual_url
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Check if it's a PDF
        content_type = response.headers.get('content-type', '').lower()
        if 'pdf' in content_type:
            return extract_text_from_pdf(response.content)
        
        # Otherwise treat as HTML
        soup = BeautifulSoup(response.content, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator="\n")
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        if not text.strip():
            raise ValueError("No content could be extracted from URL")
        
        return text
        
    except Exception as e:
        raise ValueError(f"Error extracting text from URL: {e}")

def check_dependencies():
    """Check which dependencies are available."""
    print("Checking dependencies:")
    print(f"PyMuPDF (fitz): {'✓' if PYMUPDF_AVAILABLE else '✗'}")
    print(f"PyPDF2: {'✓' if PYPDF2_AVAILABLE else '✗'}")
    print(f"pdfplumber: {'✓' if PDFPLUMBER_AVAILABLE else '✗'}")
    print(f"OCR (pytesseract + PIL): {'✓' if OCR_AVAILABLE else '✗'}")
    
    if not any([PYMUPDF_AVAILABLE, PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE]):
        print("\n❌ No PDF processing library found. Install one of:")
        print("pip install PyMuPDF  # Recommended")
        print("pip install PyPDF2")
        print("pip install pdfplumber")
    else:
        print("\n✅ At least one PDF library is available")

if __name__ == "__main__":
    check_dependencies()